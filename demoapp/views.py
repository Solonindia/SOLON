from django.contrib.auth import login, authenticate
from django.shortcuts import render, redirect
from django.contrib import messages
from django.views.decorators.csrf import csrf_protect
from .forms import CustomUserCreationForm  # Make sure LoginForm is imported from the correct path


def home_page(request):
    return render(request, 'home_page.html')

def redirect_to_home(request):
    return redirect('home')

def login_view(request):
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            return redirect('user')  # Redirect to the success page
        else:
            return render(request, 'user_login.html', {'error_message': 'Invalid credentials'})
    return render(request, 'user_login.html')

def login1_view(request):
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            #return render(request, 'success.html', {'username': username})
            return redirect('admin')
        else:
            return render(request, 'admin_login.html', {'error_message': 'Invalid credentials'})
    return render(request, 'admin_login.html')

def signup_view(request):
    if request.method == 'POST':
        form = CustomUserCreationForm(request.POST)
        if form.is_valid():
            form.save()
            return render(request, 'register.html', {'success_message': 'User created successfully'})
    else:
        form = CustomUserCreationForm()
    return render(request, 'register.html', {'form': form})

def admin_page(request):
    return render(request,'admin.html')

def user_page(request):
    return render(request,'user.html')
    
def complaint_form(request):
    return render(request,'complaint.html')

import pandas as pd
from django.shortcuts import render
from django.http import HttpResponse
from docx import Document
from io import BytesIO
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.shared import Pt

import matplotlib.pyplot as plt
import numpy as np
import base64
from io import BytesIO
import math

def generate_plot(input_data, individual_numbers, individual_POAI):
    fig, ax1 = plt.subplots()
    color_light_blue = '#0000FF'  # Hexadecimal value for light blue
    color_black = 'black'

    ax1.set_xlabel('Month')
    ax1.set_ylabel('Energy in (kWh)', color=color_black)

    bar_width = 0.95  # Adjust the width of the bars
    bar_positions = np.arange(len(input_data)) * 2  # Add space between bars

    ax1.bar(bar_positions, individual_numbers, width=bar_width, color=color_light_blue)  # Using adjusted positions and width
    ax1.tick_params(axis='y', labelcolor=color_black)

    # Calculate the upper limit for y-axis with custom intervals
    max_individual_numbers = max(individual_numbers)
    upper_limit_numbers = ((max_individual_numbers + 200000) // 200000) * 200000

    # Set custom range for y-axis
    ax1.set_ylim(0, upper_limit_numbers)

    # Set custom y-axis intervals
    yticks_numbers = np.arange(0, upper_limit_numbers + 1, 200000)
    ax1.set_yticks(yticks_numbers)

    # Format y-axis tick labels
    ax1.set_yticklabels([f'{val:,}' for val in yticks_numbers])

    ax2 = ax1.twinx()
    ax2.set_ylabel('Insolation (kWh/m2)', color=color_black)
    ax2.plot(bar_positions, individual_POAI, color='red', marker='o', linestyle='-')
    ax2.tick_params(axis='y', labelcolor=color_black)

    # Set custom range for y-axis on both axes
    upper_limit_POAI = round(max(individual_POAI) + 50)
    ax2.set_ylim(0, upper_limit_POAI)

    # Set custom y-axis intervals for ax2
    yticks_POAI = np.arange(0, upper_limit_POAI + 1, 50)
    ax2.set_yticks(yticks_POAI)

    fig.tight_layout()  # To prevent the labels from overlapping

    # Save the plot to a buffer
    buffer = BytesIO()
    plt.savefig(buffer, format='png')
    buffer.seek(0)
    image_png = buffer.getvalue()
    buffer.close()

    # Encode the plot image to base64
    graph = base64.b64encode(image_png)
    graph = graph.decode('utf-8')

    # Return both the plot image and the plot itself
    return graph,fig

def add_header(document, image_path):
    section = document.sections[0]
    header = section.header

    # Create a paragraph to control spacing before header content
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    paragraph.space_before = Pt(10)  # Adjust this value as needed for spacing

    # Create a table in the header with one row and one column
    table = header.add_table(rows=1, cols=1,width=section.page_width)
    table.autofit = True  # Disable table autofit
    table.allow_autofit = True  # Disable table autofit

    # Set the width of the first column to match the page width
    table.columns[0].width = section.page_width

    # Get the cell in the table
    cell = table.cell(0, 0)

    # Set the background color of the table cell
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '003366')  # Hex code for navy blue color
    cell._element.get_or_add_tcPr().append(shading_elm)

    # Add the logo and header text in the same cell
    paragraph_in_cell = cell.paragraphs[0]
    run = paragraph_in_cell.add_run()
    run.add_picture(image_path, width=Pt(60))
    header.add_paragraph()

def set_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tcBorders.append(border)
            tcPr.append(tcBorders)

def set_cell_margins(cell, **kwargs):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')

    for k, v in kwargs.items():
        node = OxmlElement(f'w:{k}')
        node.set(qn('w:w'), str(v))
        tcMar.append(node)
    
    tcPr.append(tcMar)

def calculate_total_time(time_list):
    # Convert each time to minutes and sum them up
    total_minutes = sum(time.hour * 60 + time.minute for time in time_list)
    # Convert total minutes to hours and remaining minutes
    total_hours, remaining_minutes = divmod(total_minutes, 60)
    # Return the total time as a tuple (hours, minutes)
    return total_hours, remaining_minutes

def Calculate_timings(new_list):
    # Calculate the total time from the new list
    total_hours = sum(item[0] for item in new_list)
    total_minutes = sum(item[1] for item in new_list)

    # Convert total hours and minutes to hours and remaining minutes
    total_hours += total_minutes // 60
    remaining_minutes = total_minutes % 60
    return str(total_hours)+':'+str(remaining_minutes)

def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "000000", "space": "0"},
        bottom={"sz": 12, "val": "single", "color": "000000", "space": "0"},
        start={"sz": 12, "val": "single", "color": "000000", "space": "0"},
        end={"sz": 12, "val": "single", "color": "000000", "space": "0"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "start", "bottom", "end", "left", "right", "insideH", "insideV"):
        if edge in kwargs:
            edge_el = tcBorders.find(qn("w:" + edge))
            if edge_el is None:
                edge_el = OxmlElement("w:" + edge)
                tcBorders.append(edge_el)
            for key in kwargs[edge]:
                edge_el.set(qn("w:" + key), str(kwargs[edge][key]))

def create_bookmark(paragraph, bookmark_text, bookmark_name):
    run = paragraph.add_run()
    tag = run._r
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), '0')
    bookmark_start.set(qn('w:name'), bookmark_name)
    tag.append(bookmark_start)

    run = paragraph.add_run(bookmark_text)

    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), '0')
    bookmark_end.set(qn('w:name'), bookmark_name)
    tag.append(bookmark_end)

def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def generate_word(request):
    if request.method == 'POST':
        try:
            # Process uploaded Excel files and extract data
            input_data = []
            individual_numbers = []
            individual_Exp = []
            individual_Imp = []
            individual_Grid = []
            individual_POAI = []
            total_sum = 0
            total_Exp = 0
            total_Imp = 0
            total_Grid = 0
            total_POAI = 0
            
            input_data1 = []
            individual_SPR = []
            individual_PRA = []
            individual_PRC = []
            individual_PLF = []
            individual_PA = []
            individual_GA = []
            total_SPR = 0
            total_PRA = 0
            total_PRC = 0
            total_PLF = 0
            total_PA = 0
            total_GA = 0

            new_smb = []
            new_grid = []
            new_others = []
            new_inv = []
            new_tranformer = []
            new_string = []
            new_loss_string = []
            new_loss_smb = []
            new_loss_grid = []
            new_loss_others = []
            new_loss_inv = []
            new_loss_tranformer = []

            for i in range(1, 13):
                uploaded_file = request.FILES.get(f'file{i}')
                input_value = request.POST.get(f'input{i}', '')
                input_data.append(input_value)
                input_data1.append(input_value)

                if uploaded_file:
                    df = pd.read_excel(uploaded_file, sheet_name='S_11.9', engine='openpyxl')
                    unnamed_22_column = df['Unnamed: 22']
                    unnamed_22_numeric = pd.to_numeric(unnamed_22_column, errors='coerce')
                    total_sum_unnamed_22_mwh = unnamed_22_numeric.sum()
                    total_sum_unnamed_22_kwh = round(total_sum_unnamed_22_mwh * 1000)
                    individual_numbers.append(total_sum_unnamed_22_kwh)
                    total_sum += total_sum_unnamed_22_kwh

                    df2 = pd.read_excel(uploaded_file, sheet_name='Plant_Start',)
                    index_of_MF = df2.columns.get_loc('MF')
                    next_column_name = df2.columns[index_of_MF + 1]
                    column_numeric = pd.to_numeric(df2[next_column_name], errors='coerce')
                    column_numeric_positive = column_numeric[column_numeric > 0]
                    sum_next_column = column_numeric_positive.sum()
                    sum_Exp = round(sum_next_column)
                    individual_Exp.append(sum_Exp)
                    total_Exp = round(total_Exp + sum_next_column)

                    df3 = pd.read_excel(uploaded_file, sheet_name='Plant_Start')
                    index_of_MF = df3.columns.get_loc('MF')
                    next_column_name = df3.columns[index_of_MF + 4]
                    column_numeric = pd.to_numeric(df3[next_column_name], errors='coerce')
                    column_numeric_positive = column_numeric[column_numeric > 0]
                    sum_next_column = column_numeric_positive.sum()
                    sum_Imp = round(sum_next_column)
                    individual_Imp.append(sum_Imp)
                    total_Imp = round(total_Imp + sum_next_column)

                    df4 = pd.read_excel(uploaded_file, sheet_name='Plant_Start')
                    index_of_MF = df3.columns.get_loc('MF')
                    next_column = df3.columns[index_of_MF + 6]
                    column_num = pd.to_numeric(df4[next_column], errors='coerce')
                    sum_data = column_num.sum()
                    sum_Grid = round(sum_data * 1000)
                    individual_Grid.append(sum_Grid)
                    total_Grid = total_Grid + sum_Grid

                    df5 = pd.read_excel(uploaded_file, sheet_name='SUMMARY')
                    cell_value = round(df5.iloc[5, 6], 2)
                    individual_POAI.append(cell_value)
                    total_POAI = round(total_POAI + cell_value, 2)

                    df1 = pd.read_excel(uploaded_file, sheet_name='SUMMARY')
                    #print(df1)
                    unnamed_7_column = df1['Unnamed: 7']
                    unnamed_5_column = df1['Unnamed: 5']
                    unnamed_24_column = df1['Unnamed: 24']
                    unnamed_7_column_0 = round(unnamed_7_column[0] * 100,2)
                    unnamed_7_column_5 = round(unnamed_7_column[5] * 100 ,2)
                    unnamed_5_column_5 = round(unnamed_5_column[5] * 100,2)
                    unnamed_24_column_5 = round(unnamed_24_column[5] * 100,2)
                    unnamed_26_column = df1['Unnamed: 26'] 
                    unnamed_17_column = df1['Unnamed: 17']
                    unnamed_26_column_5 = round(unnamed_26_column[5] * 100,2)
                    unnamed_17_column_3 = round(unnamed_17_column[3] * 100,2)
                    if math.isnan(unnamed_26_column_5):
                        individual_GA.append(unnamed_17_column_3)
                        total_GA = total_GA + unnamed_17_column_3
                    else:
                        individual_GA.append(unnamed_26_column_5)
                        total_GA = total_GA + unnamed_26_column_5
                
                    individual_SPR.append(unnamed_7_column_0)
                    individual_PRA.append(unnamed_7_column_5)
                    individual_PLF.append(unnamed_5_column_5)
                    individual_PA.append(unnamed_24_column_5)


                    total_SPR = total_SPR + unnamed_7_column_0
                    total_PRA = total_PRA + unnamed_7_column_5
                    total_PLF = total_PLF + unnamed_5_column_5
                    total_PA = total_PA + unnamed_24_column_5
                if 'Unnamed: 30' in df1.columns and 'Unnamed: 31' in df1.columns:
                    unnamed_30_column = df1['Unnamed: 30'] 
                    unnamed_31_column = df1['Unnamed: 31']
                    unnamed_30_column_5 = round(unnamed_30_column[5] * 100, 2) if not pd.isnull(unnamed_30_column[5]) else None
                    unnamed_31_column_5 = round(unnamed_31_column[5] * 100, 2) if not pd.isnull(unnamed_31_column[5]) else None
                    if pd.isnull(unnamed_30_column_5) and unnamed_31_column_5 is None: 
                       unnamed_25_column = df1['Unnamed: 25']
                       unnamed_25_column_5 = round(unnamed_25_column[5] * 100, 2)
                       individual_PRC.append(unnamed_25_column_5)
                       total_PRC = total_PRC + unnamed_25_column_5
                    elif unnamed_31_column_5 is None:  
                       individual_PRC.append(unnamed_30_column_5)
                       total_PRC = total_PRC + unnamed_30_column_5
                    else:
                       individual_PRC.append(unnamed_31_column_5)
                       total_PRC = total_PRC + unnamed_31_column_5
                    avg_prc = round(total_PRC / 12, 2)
                else:
                    if 'Unnamed: 25' in df1.columns:
                       unnamed_25_column = df1['Unnamed: 25']
                       unnamed_25_column_5 = round(unnamed_25_column[5] * 100, 2)
                       individual_PRC.append(unnamed_25_column_5)
                       total_PRC = total_PRC + unnamed_25_column_5
                    else:
                       pass
                    avg_prc = round(total_PRC / 12, 2)
                df0 = pd.read_excel(uploaded_file, sheet_name='LOSS GEN')
                filtered_df = df0[df0['Unnamed: 2'] == 'BD_SMB']
                filtered_df1 = df0[df0['Unnamed: 2'] == 'BD_Grid']
                filtered_df2 = df0[df0['Unnamed: 2'] == 'BD_Others']
                filtered_df3 = df0[df0['Unnamed: 2'] == 'BD_INV']
                filtered_df4 = df0[df0['Unnamed: 2'] == 'BD_Transformer']
                filtered_df5 = df0[df0['Unnamed: 2'] == 'BD_String']

                unnamed_10_column = df0.columns.get_loc('Unnamed: 10')
                next_column_name = df0.columns[unnamed_10_column+1]
                next_column_smb = filtered_df[next_column_name].tolist()
                next_column_grid = filtered_df1[next_column_name].tolist()
                next_column_others = filtered_df2[next_column_name].tolist()
                next_column_inv = filtered_df3[next_column_name].tolist()
                next_column_tranformer = filtered_df4[next_column_name].tolist()
                next_column_string = filtered_df5[next_column_name].tolist()

                unnamed_15_column = df0.columns.get_loc('Unnamed: 15')
                next_column = df0.columns[unnamed_15_column+1]
                next_loss_string = filtered_df5[next_column].tolist()
                next_loss_smb = filtered_df[next_column].tolist()
                next_loss_grid = filtered_df1[next_column].tolist()
                next_loss_others = filtered_df2[next_column].tolist()
                next_loss_inv = filtered_df3[next_column].tolist()
                next_loss_tranformer = filtered_df4[next_column].tolist()

                total_loss_string = sum(next_loss_string)
                total_loss_smb = sum(next_loss_smb)
                total_loss_grid = sum(next_loss_grid)
                total_loss_others = sum(next_loss_others)
                total_loss_inv = sum(next_loss_inv)
                total_loss_tranformer = sum(next_loss_tranformer)
                new_loss_string.append(total_loss_string)
                new_loss_smb.append(total_loss_smb)
                new_loss_grid.append(total_loss_grid)
                new_loss_others.append(total_loss_others)
                new_loss_inv.append(total_loss_inv)
                new_loss_tranformer.append(total_loss_tranformer)

                total_time_inv = calculate_total_time(next_column_inv)
                total_time_smb = calculate_total_time(next_column_smb)
                total_time_others = calculate_total_time(next_column_others)
                total_time_grid = calculate_total_time(next_column_grid)
                total_time_tranformer = calculate_total_time(next_column_tranformer)
                total_time_string = calculate_total_time(next_column_string)
                
                new_inv.append(total_time_inv)
                new_smb.append(total_time_smb)
                new_others.append(total_time_others)
                new_grid.append(total_time_grid)
                new_tranformer.append(total_time_tranformer)
                new_string.append(total_time_string)
            avg_spr = round(total_SPR / 12,2)
            avg_pra = round(total_PRA / 12,2)
            avg_prc = round(total_PRC / 12,2)
            avg_plf = round(total_PLF / 12,2)
            avg_pa = round(total_PA / 12,2)
            avg_ga = round(total_GA / 12, 2)
     
            total_inv = Calculate_timings(new_inv)
            total_smb = Calculate_timings(new_smb)
            total_others = Calculate_timings(new_others)
            total_grid = Calculate_timings(new_grid)
            total_tranformer = Calculate_timings(new_tranformer)
            total_string = Calculate_timings(new_string)

            total_loss_string = round(sum(new_loss_string),2)
            total_loss_smb = round(sum(new_loss_smb),2)
            total_loss_grid = round(sum(new_loss_grid),2)
            total_loss_others = round(sum(new_loss_others),2)
            total_loss_inv = round(sum(new_loss_inv),2)
            total_loss_tranformer = round(sum(new_loss_tranformer),2)

            # Generate plot
            graph, fig = generate_plot(input_data, individual_numbers, individual_POAI)
            
            # Create a new Document
            doc = Document()

            # Add header and footer
            image_path = r'static/Images/logo.jpg'
            add_header(doc, image_path)

            section = doc.sections[0]
            footer = section.footer
            footer_paragraph = footer.paragraphs[0]
            footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            add_page_number(footer_paragraph)
            # Add images
            image_path1 = r'static/Images/solar.jpg'
            doc.add_picture(image_path1, width=Inches(6.0), height=Inches(4.0))

            # Add title and content
            table = doc.add_table(rows=1, cols=1)
            table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Set the borders for the single cell (the box)
            cell = table.cell(0, 0)
            set_cell_border(
                cell,
                top={"sz": 12, "val": "single", "color": "000000"},
                bottom={"sz": 12, "val": "single", "color": "000000"},
                left={"sz": 12, "val": "single", "color": "000000"},
                right={"sz": 12, "val": "single", "color": "000000"},
            )
            # Add the paragraphs to the single cell of the table
            cell_paragraphs = [
                "Annual Performance Report",
                "FY-2022-23",
                "For CapSol Energy Private Limited",
                "11.96MWp GMFT Solar PV Power Plant,",
                "Madurai, Tamilnadu-625703"
            ]
            for text in cell_paragraphs:
                p = cell.add_paragraph(text, style='BodyText')
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            doc.add_paragraph().add_run().add_break()

            doc.add_paragraph("Submitted By", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_paragraph("SOLON India Private Limited", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_paragraph("Hyderabad-500055", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_paragraph("Date of Submission: 31.05.2022", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_section(WD_SECTION_START.NEW_PAGE)

            # Add table of contents
            heading = doc.add_paragraph("Table of Contents", style='Heading1')
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Set the font color to blue
            run = heading.runs[0]
            run.font.color.rgb = RGBColor(0, 0, 255)  # RGB for blue color            
            doc.add_paragraph().add_run().add_break()
            toc_sections = [
                "1. Plant General Information..................................................................................................3",

                "2. Energy Generation Details..................................................................................................3",
                
                "2.1. Energy Generation by Inverter.....................................................................................3",
                "2.2. Energy Exported to Grid..................................................................................................3",
                "2.3. Energy Imported from Grid............................................................................................3",
                "2.4. Irradiation POAI...................................................................................................................3",
                "2.5. Inverter Energy Vs. Insolation........................................................................................3",
                "3. Performance Analysis.............................................................................................................5",
                "3.1. Simulated Performance Ratio..........................................................................................5",
                "3.2. Actual Performance Ratio..................................................................................................5",
                "3.3. Contractual Performance Ratio......................................................................................5",
                "3.4. PLF...............................................................................................................................................5",
                "3.5. Grid Availability.....................................................................................................................5",
                "3.6. Plant Availability....................................................................................................................5",
                "4. Plant Breakdown and Downtime........................................................................................6"
            ]
            for section in toc_sections:
               doc.add_paragraph(section, style='BodyText')

            # Add a page break to start a new section for the table
            doc.add_section(WD_SECTION_START.NEW_PAGE)

            # Add table for the input data
            doc.add_paragraph("Plant General Information", style='Heading1')
            doc.add_paragraph(
                "The Solar Power plant is situated in Modagam Village, Peraiyur Taluk, Madurai District, Tamilnadu. "
                "This power plant was commissioned on 26.06.2019 and connected to Kallupatti 110/33kV Substation which is around 9 km away from Plant. "
                "There are a total of 36,600 No’s Multi-crystalline SPV Module, 915MMS Table, 78 No’s SMB, 6 No’s Inverters, and 2 No’s Inverter Duty Transformers installed at Site.",
                style='BodyText'
            )

            image_path2 = r'static/Images/topv.jpg'
            doc.add_picture(image_path2, width=Inches(6.0), height=Inches(3.0))
            doc.add_paragraph("Figure 1. Plant View", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Add sections
            sections = [
                "2. Energy Generation Details",
                "2.1. Energy Generation by Inverter",
                "2.2. Energy Exported to Grid",
                "2.3. Energy Imported from Grid",
                "2.4. Irradiation POAI",
                "2.5. Inverter Energy Vs. Insolation"
            ]
            
            for section in sections:
                doc.add_paragraph(section, style='Heading2')

            # Add a table for the input data
            doc.add_section(WD_SECTION_START.NEW_PAGE)
            table = doc.add_table(rows=1, cols=6)

            # Add table headers
            headers = ["Month", "Gen. (kWh)", "Exp. (kWh)", "Imp. (kWh)", "Grid (kWh)", "POAI (kWh/m2)"]
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True

            # Add table data
            for i in range(12):
                row_cells = table.add_row().cells
                row_cells[0].text = input_data[i]
                row_cells[1].text = str(individual_numbers[i])
                row_cells[2].text = str(individual_Exp[i])
                row_cells[3].text = str(individual_Imp[i])
                row_cells[4].text = str(individual_Grid[i])
                row_cells[5].text = str(individual_POAI[i])
                for cell in row_cells:
                    set_cell_margins(cell, top=50, start=50, bottom=50, end=50)
                    # Adjust paragraph spacing within each cell
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.space_after = Pt(0)
                        paragraph.paragraph_format.space_before = Pt(0)
                
            # Add a row for totals
            total_row = table.add_row().cells
            total_row[0].text = "Total"
            total_row[1].text = str(total_sum)
            total_row[2].text = str(total_Exp)
            total_row[3].text = str(total_Imp)
            total_row[4].text = str(total_Grid)
            total_row[5].text = str(total_POAI)

            # Make the total row bold
            for cell in total_row:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.space_before = Pt(0)
                    for run in paragraph.runs:
                        run.font.bold = True

            set_table_borders(table)
            paragraph = doc.add_paragraph()
            add_plot(doc, graph, fig)
            sections1 = [
                "3. Performance Analysis",
                "Below Table contains following information:",
                "3.1. Simulated Performance Ratio",
                "3.2. Actual Performance Ratio",
                "3.3. Contractual Performance Ratio",
                "3.4. PLF",
                "3.5. Grid Availability",
                "3.6. Plant Availability"
            ]
            
            line_spacing = Pt(9)  # Set the desired line spacing

            for section in sections1:
                para = doc.add_paragraph(section, style='Heading3')
                para_format = para.paragraph_format
                para_format.space_after = Pt(2)  # Reduce space after each paragraph
                para_format.line_spacing = line_spacing 

            table = doc.add_table(rows=1, cols=7)
            headers = ["Month", "Simulated PR%", "Actual PR%", "Corrected PR%", "PLF%", "Grid Availability%", "Plant Availability%"]
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                # Set cell margins to reduce spacing
                set_cell_margins(hdr_cells[i], top=50, start=50, bottom=50, end=50)
            # Add table data
            for i in range(12):
                row_cells = table.add_row().cells
                row_cells[0].text = input_data1[i]
                row_cells[1].text = str(individual_SPR[i])
                row_cells[2].text = str(individual_PRA[i])
                row_cells[3].text = str(individual_PRC[i])
                row_cells[4].text = str(individual_PLF[i])
                row_cells[5].text = str(individual_GA[i])
                row_cells[6].text = str(individual_PA[i])
                # Set cell margins to reduce spacing
                for cell in row_cells:
                    set_cell_margins(cell, top=50, start=50, bottom=50, end=50)
                    # Adjust paragraph spacing within each cell
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.space_after = Pt(0)
                        paragraph.paragraph_format.space_before = Pt(0)
                
            # Add a row for totals
            total_row = table.add_row().cells
            total_row[0].text = "Total"
            total_row[1].text = str(avg_spr)
            total_row[2].text = str(avg_pra)
            total_row[3].text = str(avg_prc)
            total_row[4].text = str(avg_plf)
            total_row[5].text = str(avg_ga)
            total_row[6].text = str(avg_pa)

            # Make the total row bold and reduce spacing
            for cell in total_row:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.space_before = Pt(0)
                    for run in paragraph.runs:
                        run.font.bold = True

            set_table_borders(table)
            sections2 = [
                "Note: ",
                "As per Amended O&M Contract Dated: 07.08.2020",
                "Annexure-B, the contractual Plant Availability should be > 99.5% and,",
                "Annexure-C, the Contractual Performance Ratio should be > 78.5%",
            ]

            # Set the desired line spacing
            line_spacing = Pt(9)

            # Add the sections to the document with reduced line spacing
            for section in sections2:
                para = doc.add_paragraph()
                run = para.add_run(section)
                run.bold = True
                font = run.font
                font.color.rgb = RGBColor(0xFF, 0xA5, 0x00)
    
                para_format = para.paragraph_format
                para_format.space_after = Pt(2)  # Reduce space after each paragraph
                para_format.line_spacing = line_spacing  
            doc.add_paragraph("Reason for Less Plant Availability:", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            doc.add_paragraph(
                "we had a breakdown in inverter No-02, ISU-1 LCL Capacitor bank and ISU-02 IGBT Module Failure on 11th Aug-2022, due to unavailability of these spare at ABB factory warehouse it took more time to get it deliver to site. We received the materials at site on 9th Sept-2022. During these period inverters was running at 55% Load. As per ABB IGBT Module delivery got delay due to shortage of semi-conductors worldwide. In our spare list also, these spares are not available",
                style='BodyText'
            )
            p1 = doc.add_paragraph("4. Plant Breakdown and Downtime", style='BodyText')
            p1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run1 = p1.runs[0]
            run1.font.size = Pt(14)  # Increase font size
            run1.font.color.rgb = RGBColor(0, 0, 255)  # Change color to blue

            # Adding and formatting the second paragraph
            p2 = doc.add_paragraph("For the Period from 01.04.2022 to 31.03.2023", style='BodyText')
            p2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run2 = p2.runs[0]
            run2.font.size = Pt(14)  # Increase font size
            run2.font.color.rgb = RGBColor(0, 0, 255) 
            # Initialize the table
            table = doc.add_table(rows=1, cols=4)  # Adjust cols as per your headers count

            # Headers
            headers = ["S.No", "Breakdown Type", "Breakdown in Hours", "Generation Loss in MWh"]
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True

            # Example data to be added directly to columns 1 and 2
            data_columns_1_2 = [
                ["A", "Controllable Events"],
                ["1", "String Breakdown", total_string, total_loss_string],
                ["2", "SMB Breakdown", total_smb, total_loss_smb],
                ["3", "Inverter Breakdown", total_inv, total_loss_inv],
                ["4", "Transformer Breakdown", total_tranformer, total_loss_tranformer],
                ["5", "Others Breakdown", total_others, total_loss_others],
                ["B", "NON - Controllable Events"],
                ["1", "Curtailment Breakdown", "0.00", "0.00"],
                ["2", "Grid Breakdown", total_grid, total_loss_grid],
                ["3", "Force Majeure*", "0.00", "0.00"],
            ]

            # Add data to the table
            for row_data in data_columns_1_2:
                row_cells = table.add_row().cells
                for idx, cell_data in enumerate(row_data):
                    cell = row_cells[idx]
                    run = cell.paragraphs[0].add_run(str(cell_data))
                    if cell_data in ["A", "B","Controllable Events","NON - Controllable Events"]:
                        run.bold = True
            # Set table borders
            set_table_borders(table)            
            # Save the document
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            # Return the document as a response
            response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename=generated_document.docx'
            return response

        except Exception as e:
            return HttpResponse(f"An error occurred: {e}")

    upload_range = range(1, 13)
    
    return render(request, 'upload_form1.html', {'upload_range': upload_range})

def add_plot(document, graph, fig):
    paragraph = document.add_paragraph()
    run_title = paragraph.add_run("Energy Generation vs. Insolation Trend")
    run_title.bold = True
    font_title = run_title.font
    font_title.color.rgb = RGBColor(0xFF, 0x69, 0xB4)  # Set font color to pink
    font_title.size = Pt(16)  # Set font size to 16 points
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_paragraph()
    run_image = document.add_paragraph().add_run()
    image_png = base64.b64decode(graph)
    image_stream = BytesIO(image_png)
    run_image.add_picture(image_stream, width=Inches(6.0), height=Inches(3.0))

def generate_word1(request):
    if request.method == 'POST':
        try:
            # Process uploaded Excel files and extract data
            input_data = []
            individual_numbers = []
            individual_Exp = []
            individual_Imp = []
            individual_Grid = []
            individual_POAI = []
            total_sum = 0
            total_Exp = 0
            total_Imp = 0
            total_Grid = 0
            total_POAI = 0
            
            input_data1 = []
            individual_SPR = []
            individual_PRA = []
            individual_PRC = []
            individual_PLF = []
            individual_PA = []
            individual_GA = []
            total_SPR = 0
            total_PRA = 0
            total_PRC = 0
            total_PLF = 0
            total_PA = 0
            total_GA = 0

            new_smb = []
            new_grid = []
            new_others = []
            new_inv = []
            new_tranformer = []
            new_string = []
            new_loss_string = []
            new_loss_smb = []
            new_loss_grid = []
            new_loss_others = []
            new_loss_inv = []
            new_loss_tranformer = []

            for i in range(1, 13):
                uploaded_file = request.FILES.get(f'file{i}')
                input_value = request.POST.get(f'input{i}', '')
                input_data.append(input_value)
                input_data1.append(input_value)

                if uploaded_file:
                    df = pd.read_excel(uploaded_file, sheet_name='S_11.9', engine='openpyxl')
                    unnamed_22_column = df['Unnamed: 22']
                    unnamed_22_numeric = pd.to_numeric(unnamed_22_column, errors='coerce')
                    total_sum_unnamed_22_mwh = unnamed_22_numeric.sum()
                    total_sum_unnamed_22_kwh = round(total_sum_unnamed_22_mwh * 1000)
                    individual_numbers.append(total_sum_unnamed_22_kwh)
                    total_sum += total_sum_unnamed_22_kwh

                    df2 = pd.read_excel(uploaded_file, sheet_name='Plant_Start',)
                    index_of_MF = df2.columns.get_loc('MF')
                    next_column_name = df2.columns[index_of_MF + 1]
                    column_numeric = pd.to_numeric(df2[next_column_name], errors='coerce')
                    column_numeric_positive = column_numeric[column_numeric > 0]
                    sum_next_column = column_numeric_positive.sum()
                    sum_Exp = round(sum_next_column)
                    individual_Exp.append(sum_Exp)
                    total_Exp = round(total_Exp + sum_next_column)

                    df3 = pd.read_excel(uploaded_file, sheet_name='Plant_Start')
                    index_of_MF = df3.columns.get_loc('MF')
                    next_column_name = df3.columns[index_of_MF + 4]
                    column_numeric = pd.to_numeric(df3[next_column_name], errors='coerce')
                    column_numeric_positive = column_numeric[column_numeric > 0]
                    sum_next_column = column_numeric_positive.sum()
                    sum_Imp = round(sum_next_column)
                    individual_Imp.append(sum_Imp)
                    total_Imp = round(total_Imp + sum_next_column)

                    df4 = pd.read_excel(uploaded_file, sheet_name='Plant_Start')
                    index_of_MF = df3.columns.get_loc('MF')
                    next_column = df3.columns[index_of_MF + 6]
                    column_num = pd.to_numeric(df4[next_column], errors='coerce')
                    sum_data = column_num.sum()
                    sum_Grid = round(sum_data * 1000)
                    individual_Grid.append(sum_Grid)
                    total_Grid = total_Grid + sum_Grid

                    df5 = pd.read_excel(uploaded_file, sheet_name='SUMMARY')
                    cell_value = round(df5.iloc[5, 6], 2)
                    individual_POAI.append(cell_value)
                    total_POAI = round(total_POAI + cell_value, 2)

                    df1 = pd.read_excel(uploaded_file, sheet_name='SUMMARY')
                    #print(df1)
                    unnamed_7_column = df1['Unnamed: 7']
                    unnamed_5_column = df1['Unnamed: 5']
                    unnamed_24_column = df1['Unnamed: 24']
                    unnamed_7_column_0 = round(unnamed_7_column[0] * 100,2)
                    unnamed_7_column_5 = round(unnamed_7_column[5] * 100 ,2)
                    unnamed_5_column_5 = round(unnamed_5_column[5] * 100,2)
                    unnamed_24_column_5 = round(unnamed_24_column[5] * 100,2)
                    unnamed_26_column = df1['Unnamed: 26'] 
                    unnamed_17_column = df1['Unnamed: 17']
                    unnamed_26_column_5 = round(unnamed_26_column[5] * 100,2)
                    unnamed_17_column_3 = round(unnamed_17_column[3] * 100,2)
                    if math.isnan(unnamed_26_column_5):
                        individual_GA.append(unnamed_17_column_3)
                        total_GA = total_GA + unnamed_17_column_3
                    else:
                        individual_GA.append(unnamed_26_column_5)
                        total_GA = total_GA + unnamed_26_column_5
                
                    individual_SPR.append(unnamed_7_column_0)
                    individual_PRA.append(unnamed_7_column_5)
                    individual_PLF.append(unnamed_5_column_5)
                    individual_PA.append(unnamed_24_column_5)


                    total_SPR = total_SPR + unnamed_7_column_0
                    total_PRA = total_PRA + unnamed_7_column_5
                    total_PLF = total_PLF + unnamed_5_column_5
                    total_PA = total_PA + unnamed_24_column_5
                if 'Unnamed: 30' in df1.columns and 'Unnamed: 31' in df1.columns:
                    unnamed_30_column = df1['Unnamed: 30'] 
                    unnamed_31_column = df1['Unnamed: 31']
                    unnamed_30_column_5 = round(unnamed_30_column[5] * 100, 2) if not pd.isnull(unnamed_30_column[5]) else None
                    unnamed_31_column_5 = round(unnamed_31_column[5] * 100, 2) if not pd.isnull(unnamed_31_column[5]) else None
                    if pd.isnull(unnamed_30_column_5) and unnamed_31_column_5 is None: 
                       unnamed_25_column = df1['Unnamed: 25']
                       unnamed_25_column_5 = round(unnamed_25_column[5] * 100, 2)
                       individual_PRC.append(unnamed_25_column_5)
                       total_PRC = total_PRC + unnamed_25_column_5
                    elif unnamed_31_column_5 is None:  
                       individual_PRC.append(unnamed_30_column_5)
                       total_PRC = total_PRC + unnamed_30_column_5
                    else:
                       individual_PRC.append(unnamed_31_column_5)
                       total_PRC = total_PRC + unnamed_31_column_5
                    avg_prc = round(total_PRC / 12, 2)
                else:
                    if 'Unnamed: 25' in df1.columns:
                       unnamed_25_column = df1['Unnamed: 25']
                       unnamed_25_column_5 = round(unnamed_25_column[5] * 100, 2)
                       individual_PRC.append(unnamed_25_column_5)
                       total_PRC = total_PRC + unnamed_25_column_5
                    else:
                       pass
                    avg_prc = round(total_PRC / 12, 2)
                df0 = pd.read_excel(uploaded_file, sheet_name='LOSS GEN')
                filtered_df = df0[df0['Unnamed: 2'] == 'BD_SMB']
                filtered_df1 = df0[df0['Unnamed: 2'] == 'BD_Grid']
                filtered_df2 = df0[df0['Unnamed: 2'] == 'BD_Others']
                filtered_df3 = df0[df0['Unnamed: 2'] == 'BD_INV']
                filtered_df4 = df0[df0['Unnamed: 2'] == 'BD_Transformer']
                filtered_df5 = df0[df0['Unnamed: 2'] == 'BD_String']

                unnamed_10_column = df0.columns.get_loc('Unnamed: 10')
                next_column_name = df0.columns[unnamed_10_column+1]
                next_column_smb = filtered_df[next_column_name].tolist()
                next_column_grid = filtered_df1[next_column_name].tolist()
                next_column_others = filtered_df2[next_column_name].tolist()
                next_column_inv = filtered_df3[next_column_name].tolist()
                next_column_tranformer = filtered_df4[next_column_name].tolist()
                next_column_string = filtered_df5[next_column_name].tolist()

                unnamed_15_column = df0.columns.get_loc('Unnamed: 15')
                next_column = df0.columns[unnamed_15_column+1]
                next_loss_string = filtered_df5[next_column].tolist()
                next_loss_smb = filtered_df[next_column].tolist()
                next_loss_grid = filtered_df1[next_column].tolist()
                next_loss_others = filtered_df2[next_column].tolist()
                next_loss_inv = filtered_df3[next_column].tolist()
                next_loss_tranformer = filtered_df4[next_column].tolist()

                total_loss_string = sum(next_loss_string)
                total_loss_smb = sum(next_loss_smb)
                total_loss_grid = sum(next_loss_grid)
                total_loss_others = sum(next_loss_others)
                total_loss_inv = sum(next_loss_inv)
                total_loss_tranformer = sum(next_loss_tranformer)
                new_loss_string.append(total_loss_string)
                new_loss_smb.append(total_loss_smb)
                new_loss_grid.append(total_loss_grid)
                new_loss_others.append(total_loss_others)
                new_loss_inv.append(total_loss_inv)
                new_loss_tranformer.append(total_loss_tranformer)

                total_time_inv = calculate_total_time(next_column_inv)
                total_time_smb = calculate_total_time(next_column_smb)
                total_time_others = calculate_total_time(next_column_others)
                total_time_grid = calculate_total_time(next_column_grid)
                total_time_tranformer = calculate_total_time(next_column_tranformer)
                total_time_string = calculate_total_time(next_column_string)
                
                new_inv.append(total_time_inv)
                new_smb.append(total_time_smb)
                new_others.append(total_time_others)
                new_grid.append(total_time_grid)
                new_tranformer.append(total_time_tranformer)
                new_string.append(total_time_string)
            avg_spr = round(total_SPR / 12,2)
            avg_pra = round(total_PRA / 12,2)
            avg_prc = round(total_PRC / 12,2)
            avg_plf = round(total_PLF / 12,2)
            avg_pa = round(total_PA / 12,2)
            avg_ga = round(total_GA / 12, 2)
     
            total_inv = Calculate_timings(new_inv)
            total_smb = Calculate_timings(new_smb)
            total_others = Calculate_timings(new_others)
            total_grid = Calculate_timings(new_grid)
            total_tranformer = Calculate_timings(new_tranformer)
            total_string = Calculate_timings(new_string)

            total_loss_string = round(sum(new_loss_string),2)
            total_loss_smb = round(sum(new_loss_smb),2)
            total_loss_grid = round(sum(new_loss_grid),2)
            total_loss_others = round(sum(new_loss_others),2)
            total_loss_inv = round(sum(new_loss_inv),2)
            total_loss_tranformer = round(sum(new_loss_tranformer),2)

            # Generate plot
            graph, fig = generate_plot(input_data, individual_numbers, individual_POAI)
            
            # Create a new Document
            doc = Document()

            # Add header and footer
            image_path = r'static/Images/logo.jpg'
            add_header(doc, image_path)

            section = doc.sections[0]
            footer = section.footer
            footer_paragraph = footer.paragraphs[0]
            footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            add_page_number(footer_paragraph)
            # Add images
            image_path1 = r'static/Images/solar.jpg'
            doc.add_picture(image_path1, width=Inches(6.0), height=Inches(4.0))

            # Add title and content
            table = doc.add_table(rows=1, cols=1)
            table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Set the borders for the single cell (the box)
            cell = table.cell(0, 0)
            set_cell_border(
                cell,
                top={"sz": 12, "val": "single", "color": "000000"},
                bottom={"sz": 12, "val": "single", "color": "000000"},
                left={"sz": 12, "val": "single", "color": "000000"},
                right={"sz": 12, "val": "single", "color": "000000"},
            )
            # Add the paragraphs to the single cell of the table
            cell_paragraphs = [
                "Annual Performance Report",
                "FY-2022-23",
                "For CapSol Energy Private Limited",
                "11.96MWp GMFT Solar PV Power Plant,",
                "Madurai, Tamilnadu-625703"
            ]
            for text in cell_paragraphs:
                p = cell.add_paragraph(text, style='BodyText')
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            doc.add_paragraph().add_run().add_break()

            doc.add_paragraph("Submitted By", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_paragraph("SOLON India Private Limited", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_paragraph("Hyderabad-500055", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_paragraph("Date of Submission: 31.05.2022", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_section(WD_SECTION_START.NEW_PAGE)

            # Add table of contents
            heading = doc.add_paragraph("Table of Contents", style='Heading1')
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Set the font color to blue
            run = heading.runs[0]
            run.font.color.rgb = RGBColor(0, 0, 255)  # RGB for blue color            
            doc.add_paragraph().add_run().add_break()
            toc_sections = [
                "1. Plant General Information..................................................................................................3",

                "2. Energy Generation Details..................................................................................................3",
                
                "2.1. Energy Generation by Inverter.....................................................................................3",
                "2.2. Energy Exported to Grid..................................................................................................3",
                "2.3. Energy Imported from Grid............................................................................................3",
                "2.4. Irradiation POAI...................................................................................................................3",
                "2.5. Inverter Energy Vs. Insolation........................................................................................3",
                "3. Performance Analysis.............................................................................................................5",
                "3.1. Simulated Performance Ratio..........................................................................................5",
                "3.2. Actual Performance Ratio..................................................................................................5",
                "3.3. Contractual Performance Ratio......................................................................................5",
                "3.4. PLF...............................................................................................................................................5",
                "3.5. Grid Availability.....................................................................................................................5",
                "3.6. Plant Availability....................................................................................................................5",
                "4. Plant Breakdown and Downtime........................................................................................6"
            ]
            for section in toc_sections:
               doc.add_paragraph(section, style='BodyText')

            # Add a page break to start a new section for the table
            doc.add_section(WD_SECTION_START.NEW_PAGE)

            # Add table for the input data
            doc.add_paragraph("Plant General Information", style='Heading1')
            doc.add_paragraph(
                "The Solar Power plant is situated in Modagam Village, Peraiyur Taluk, Madurai District, Tamilnadu. "
                "This power plant was commissioned on 26.06.2019 and connected to Kallupatti 110/33kV Substation which is around 9 km away from Plant. "
                "There are a total of 36,600 No’s Multi-crystalline SPV Module, 915MMS Table, 78 No’s SMB, 6 No’s Inverters, and 2 No’s Inverter Duty Transformers installed at Site.",
                style='BodyText'
            )

            image_path2 = r'static/Images/topv.jpg'
            doc.add_picture(image_path2, width=Inches(6.0), height=Inches(3.0))
            doc.add_paragraph("Figure 1. Plant View", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Add sections
            sections = [
                "2. Energy Generation Details",
                "2.1. Energy Generation by Inverter",
                "2.2. Energy Exported to Grid",
                "2.3. Energy Imported from Grid",
                "2.4. Irradiation POAI",
                "2.5. Inverter Energy Vs. Insolation"
            ]
            
            for section in sections:
                doc.add_paragraph(section, style='Heading2')

            # Add a table for the input data
            doc.add_section(WD_SECTION_START.NEW_PAGE)
            table = doc.add_table(rows=1, cols=6)

            # Add table headers
            headers = ["Month", "Gen. (kWh)", "Exp. (kWh)", "Imp. (kWh)", "Grid (kWh)", "POAI (kWh/m2)"]
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True

            # Add table data
            for i in range(12):
                row_cells = table.add_row().cells
                row_cells[0].text = input_data[i]
                row_cells[1].text = str(individual_numbers[i])
                row_cells[2].text = str(individual_Exp[i])
                row_cells[3].text = str(individual_Imp[i])
                row_cells[4].text = str(individual_Grid[i])
                row_cells[5].text = str(individual_POAI[i])
                for cell in row_cells:
                    set_cell_margins(cell, top=50, start=50, bottom=50, end=50)
                    # Adjust paragraph spacing within each cell
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.space_after = Pt(0)
                        paragraph.paragraph_format.space_before = Pt(0)
                
            # Add a row for totals
            total_row = table.add_row().cells
            total_row[0].text = "Total"
            total_row[1].text = str(total_sum)
            total_row[2].text = str(total_Exp)
            total_row[3].text = str(total_Imp)
            total_row[4].text = str(total_Grid)
            total_row[5].text = str(total_POAI)

            # Make the total row bold
            for cell in total_row:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.space_before = Pt(0)
                    for run in paragraph.runs:
                        run.font.bold = True

            set_table_borders(table)
            paragraph = doc.add_paragraph()
            add_plot(doc, graph, fig)
            sections1 = [
                "3. Performance Analysis",
                "Below Table contains following information:",
                "3.1. Simulated Performance Ratio",
                "3.2. Actual Performance Ratio",
                "3.3. Contractual Performance Ratio",
                "3.4. PLF",
                "3.5. Grid Availability",
                "3.6. Plant Availability"
            ]
            
            line_spacing = Pt(9)  # Set the desired line spacing

            for section in sections1:
                para = doc.add_paragraph(section, style='Heading3')
                para_format = para.paragraph_format
                para_format.space_after = Pt(2)  # Reduce space after each paragraph
                para_format.line_spacing = line_spacing 

            table = doc.add_table(rows=1, cols=7)
            headers = ["Month", "Simulated PR%", "Actual PR%", "Corrected PR%", "PLF%", "Grid Availability%", "Plant Availability%"]
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                # Set cell margins to reduce spacing
                set_cell_margins(hdr_cells[i], top=50, start=50, bottom=50, end=50)
            # Add table data
            for i in range(12):
                row_cells = table.add_row().cells
                row_cells[0].text = input_data1[i]
                row_cells[1].text = str(individual_SPR[i])
                row_cells[2].text = str(individual_PRA[i])
                row_cells[3].text = str(individual_PRC[i])
                row_cells[4].text = str(individual_PLF[i])
                row_cells[5].text = str(individual_GA[i])
                row_cells[6].text = str(individual_PA[i])
                # Set cell margins to reduce spacing
                for cell in row_cells:
                    set_cell_margins(cell, top=50, start=50, bottom=50, end=50)
                    # Adjust paragraph spacing within each cell
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.space_after = Pt(0)
                        paragraph.paragraph_format.space_before = Pt(0)
                
            # Add a row for totals
            total_row = table.add_row().cells
            total_row[0].text = "Total"
            total_row[1].text = str(avg_spr)
            total_row[2].text = str(avg_pra)
            total_row[3].text = str(avg_prc)
            total_row[4].text = str(avg_plf)
            total_row[5].text = str(avg_ga)
            total_row[6].text = str(avg_pa)

            # Make the total row bold and reduce spacing
            for cell in total_row:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.space_before = Pt(0)
                    for run in paragraph.runs:
                        run.font.bold = True

            set_table_borders(table)
            sections2 = [
                "Note: ",
                "As per Amended O&M Contract Dated: 07.08.2020",
                "Annexure-B, the contractual Plant Availability should be > 99.5% and,",
                "Annexure-C, the Contractual Performance Ratio should be > 78.5%",
            ]

            # Set the desired line spacing
            line_spacing = Pt(9)

            # Add the sections to the document with reduced line spacing
            for section in sections2:
                para = doc.add_paragraph()
                run = para.add_run(section)
                run.bold = True
                font = run.font
                font.color.rgb = RGBColor(0xFF, 0xA5, 0x00)
    
                para_format = para.paragraph_format
                para_format.space_after = Pt(2)  # Reduce space after each paragraph
                para_format.line_spacing = line_spacing  
            doc.add_paragraph("Reason for Less Plant Availability:", style='BodyText').alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            doc.add_paragraph(
                "we had a breakdown in inverter No-02, ISU-1 LCL Capacitor bank and ISU-02 IGBT Module Failure on 11th Aug-2022, due to unavailability of these spare at ABB factory warehouse it took more time to get it deliver to site. We received the materials at site on 9th Sept-2022. During these period inverters was running at 55% Load. As per ABB IGBT Module delivery got delay due to shortage of semi-conductors worldwide. In our spare list also, these spares are not available",
                style='BodyText'
            )
            p1 = doc.add_paragraph("4. Plant Breakdown and Downtime", style='BodyText')
            p1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run1 = p1.runs[0]
            run1.font.size = Pt(14)  # Increase font size
            run1.font.color.rgb = RGBColor(0, 0, 255)  # Change color to blue

            # Adding and formatting the second paragraph
            p2 = doc.add_paragraph("For the Period from 01.04.2022 to 31.03.2023", style='BodyText')
            p2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run2 = p2.runs[0]
            run2.font.size = Pt(14)  # Increase font size
            run2.font.color.rgb = RGBColor(0, 0, 255) 
            # Initialize the table
            table = doc.add_table(rows=1, cols=4)  # Adjust cols as per your headers count

            # Headers
            headers = ["S.No", "Breakdown Type", "Breakdown in Hours", "Generation Loss in MWh"]
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True

            # Example data to be added directly to columns 1 and 2
            data_columns_1_2 = [
                ["A", "Controllable Events"],
                ["1", "String Breakdown", total_string, total_loss_string],
                ["2", "SMB Breakdown", total_smb, total_loss_smb],
                ["3", "Inverter Breakdown", total_inv, total_loss_inv],
                ["4", "Transformer Breakdown", total_tranformer, total_loss_tranformer],
                ["5", "Others Breakdown", total_others, total_loss_others],
                ["B", "NON - Controllable Events"],
                ["1", "Curtailment Breakdown", "0.00", "0.00"],
                ["2", "Grid Breakdown", total_grid, total_loss_grid],
                ["3", "Force Majeure*", "0.00", "0.00"],
            ]

            # Add data to the table
            for row_data in data_columns_1_2:
                row_cells = table.add_row().cells
                for idx, cell_data in enumerate(row_data):
                    cell = row_cells[idx]
                    run = cell.paragraphs[0].add_run(str(cell_data))
                    if cell_data in ["A", "B","Controllable Events","NON - Controllable Events"]:
                        run.bold = True
            # Set table borders
            set_table_borders(table)            
            # Save the document
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            # Return the document as a response
            response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename=generated_document.docx'
            return response

        except Exception as e:
            return HttpResponse(f"An error occurred: {e}")

    upload_range = range(1, 13)
    
    return render(request, 'upload_form.html', {'upload_range': upload_range})

def add_plot(document, graph, fig):
    paragraph = document.add_paragraph()
    run_title = paragraph.add_run("Energy Generation vs. Insolation Trend")
    run_title.bold = True
    font_title = run_title.font
    font_title.color.rgb = RGBColor(0xFF, 0x69, 0xB4)  # Set font color to pink
    font_title.size = Pt(16)  # Set font size to 16 points
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_paragraph()
    run_image = document.add_paragraph().add_run()
    image_png = base64.b64decode(graph)
    image_stream = BytesIO(image_png)
    run_image.add_picture(image_stream, width=Inches(6.0), height=Inches(3.0))
